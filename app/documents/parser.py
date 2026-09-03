import csv
import io
from pathlib import Path
import docx
import pypdf
from fastapi import HTTPException
from app.core.config import settings


def parse_pdf(content_bytes: bytes) -> str:
    """Extracts text from a PDF file in memory. Scanned PDFs with no text layer raise 422."""
    try:
        reader = pypdf.PdfReader(io.BytesIO(content_bytes))
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception:
                raise HTTPException(
                    status_code=422,
                    detail="Encrypted PDF files are not supported.",
                )

        pages_text = []
        for page in reader.pages:
            text = page.extract_text() or ""
            if text.strip():
                pages_text.append(text.strip())

        extracted = "\n\n".join(pages_text).strip()
        if not extracted:
            raise HTTPException(
                status_code=422,
                detail="The PDF contains no extractable text. Scanned PDFs without an embedded text layer are not supported.",
            )
        return extracted
    except HTTPException:
        raise
    except Exception:
        raise HTTPException(
            status_code=422,
            detail="Failed to parse PDF document: invalid or corrupted file.",
        )


def parse_docx(content_bytes: bytes) -> str:
    """Extracts paragraphs and tables from a Word (.docx) document in memory."""
    try:
        doc = docx.Document(io.BytesIO(content_bytes))
        elements = []

        for p in doc.paragraphs:
            if p.text.strip():
                elements.append(p.text.strip())

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    elements.append(" | ".join(cells))

        extracted = "\n\n".join(elements).strip()
        if not extracted:
            raise HTTPException(
                status_code=422,
                detail="The Word document (.docx) contains no extractable text.",
            )
        return extracted
    except HTTPException:
        raise
    except Exception:
        raise HTTPException(
            status_code=422,
            detail="Failed to parse Word document: invalid or corrupted file.",
        )


def parse_txt_or_md(content_bytes: bytes) -> str:
    """Decodes plain text or Markdown files."""
    try:
        text = content_bytes.decode("utf-8")
    except UnicodeDecodeError:
        try:
            text = content_bytes.decode("latin-1")
        except Exception:
            raise HTTPException(
                status_code=422,
                detail="Unable to decode text file. Ensure it is UTF-8 or standard text encoding.",
            )

    extracted = text.strip()
    if not extracted:
        raise HTTPException(
            status_code=422,
            detail="The text file is empty or contains only whitespace.",
        )
    return extracted


def parse_csv(content_bytes: bytes) -> str:
    """Converts CSV file into a clean textual representation."""
    try:
        text = content_bytes.decode("utf-8")
    except UnicodeDecodeError:
        try:
            text = content_bytes.decode("latin-1")
        except Exception:
            raise HTTPException(
                status_code=422,
                detail="Unable to decode CSV file. Ensure it is UTF-8 or standard text encoding.",
            )

    try:
        sample = text[:2048]
        try:
            dialect = csv.Sniffer().sniff(sample)
            reader = csv.reader(io.StringIO(text), dialect)
        except Exception:
            reader = csv.reader(io.StringIO(text))

        rows = list(reader)
    except Exception:
        raise HTTPException(
            status_code=422,
            detail="Failed to parse CSV document.",
        )

    formatted_rows = []
    for row in rows:
        cleaned_cells = [cell.strip() for cell in row]
        if any(cleaned_cells):
            formatted_rows.append(" | ".join(cleaned_cells))

    extracted = "\n".join(formatted_rows).strip()
    if not extracted:
        raise HTTPException(
            status_code=422,
            detail="The CSV file contains no data rows.",
        )
    return extracted


def parse_document(filename: str, content_bytes: bytes) -> str:
    """
    Parses document content into text based on file extension.
    Applies truncation if the text exceeds DOCUMENT_MAX_CHARS.
    """
    suffix = Path(filename).suffix.lower()

    if suffix == ".pdf":
        raw_text = parse_pdf(content_bytes)
    elif suffix == ".docx":
        raw_text = parse_docx(content_bytes)
    elif suffix in {".txt", ".md"}:
        raw_text = parse_txt_or_md(content_bytes)
    elif suffix == ".csv":
        raw_text = parse_csv(content_bytes)
    else:
        raise HTTPException(
            status_code=415,
            detail=f"Unsupported file format '{suffix}'. Supported formats: .pdf, .docx, .txt, .md, .csv",
        )

    max_chars = settings.document_max_chars
    if len(raw_text) > max_chars:
        raw_text = (
            f"{raw_text[:max_chars]}\n\n"
            f"[NOTA INTERNA: El documento superó el límite de {max_chars} caracteres "
            f"y fue truncado de forma controlada. Se procesaron los primeros {max_chars} de {len(raw_text)} caracteres.]"
        )

    return raw_text
