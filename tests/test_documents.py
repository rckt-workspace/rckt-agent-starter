import io
from unittest.mock import AsyncMock, patch
import pytest
from fastapi.testclient import TestClient
import docx
import pypdf

from app.main import app
from app.core.config import settings
from app.agent.agent import Agent
from app.documents.parser import parse_document


@pytest.fixture
def client():
    return TestClient(app)


@pytest.fixture
def mock_openrouter():
    """Mock OpenRouterClient.complete_with_system to prevent real API calls."""
    with patch(
        "app.agent.agent.OpenRouterClient.complete_with_system",
        new_callable=AsyncMock,
    ) as mock_complete:
        mock_complete.return_value = "Esta es una respuesta simulada del modelo."
        yield mock_complete


def create_in_memory_pdf(text: str = "Texto de prueba en documento PDF.") -> bytes:
    """Generates a minimal valid PDF with an embedded text stream."""
    stream_content = f"BT /F1 12 Tf 50 250 Td ({text}) Tj ET\n".encode("latin-1")
    length = len(stream_content)
    header = (
        b"%PDF-1.4\n"
        b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n"
        b"2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj\n"
        b"3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 300 300] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >> endobj\n"
    )
    stream_obj = f"4 0 obj << /Length {length} >>\nstream\n".encode("ascii") + stream_content + b"endstream\nendobj\n"
    footer = (
        b"5 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj\n"
        b"xref\n0 6\n0000000000 65535 f \n"
        b"trailer << /Size 6 /Root 1 0 R >>\nstartxref\n300\n%%EOF"
    )
    return header + stream_obj + footer


def create_in_memory_empty_pdf() -> bytes:
    """Generates a valid PDF with a blank page and no text layer."""
    writer = pypdf.PdfWriter()
    writer.add_blank_page(width=100, height=100)
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def create_in_memory_docx(text: str = "Texto de prueba en documento DOCX.") -> bytes:
    """Generates a valid DOCX document in memory."""
    doc = docx.Document()
    doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def create_in_memory_empty_docx() -> bytes:
    """Generates a valid but empty DOCX document."""
    doc = docx.Document()
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# 1. TXT válido
# ---------------------------------------------------------------------------
def test_chat_document_valid_txt(client, mock_openrouter):
    txt_content = b"Este es el contenido de un archivo de texto plano para testing."
    files = {"file": ("nota.txt", txt_content, "text/plain")}
    data = {"message": "Resume este archivo"}

    response = client.post("/api/chat/document", data=data, files=files)

    assert response.status_code == 200
    json_data = response.json()
    assert "answer" in json_data
    assert json_data["answer"] == "Esta es una respuesta simulada del modelo."


# ---------------------------------------------------------------------------
# 2. Markdown válido
# ---------------------------------------------------------------------------
def test_chat_document_valid_markdown(client, mock_openrouter):
    md_content = b"# Guia del Proyecto\n\nEste es un archivo **markdown** con especificaciones."
    files = {"file": ("guia.md", md_content, "text/markdown")}
    data = {"message": "Explica la guia"}

    response = client.post("/api/chat/document", data=data, files=files)

    assert response.status_code == 200
    assert response.json()["answer"] == "Esta es una respuesta simulada del modelo."


# ---------------------------------------------------------------------------
# 3. CSV válido
# ---------------------------------------------------------------------------
def test_chat_document_valid_csv(client, mock_openrouter):
    csv_content = b"nombre,puesto,departamento\nAna,Lider,Ingenieria\nLuis,Disenador,Producto"
    files = {"file": ("equipo.csv", csv_content, "text/csv")}
    data = {"message": "Quien es la lider?"}

    response = client.post("/api/chat/document", data=data, files=files)

    assert response.status_code == 200
    assert response.json()["answer"] == "Esta es una respuesta simulada del modelo."


# ---------------------------------------------------------------------------
# 4. Formato no soportado (415)
# ---------------------------------------------------------------------------
def test_chat_document_unsupported_format(client):
    exe_content = b"MZ\x90\x00\x03\x00\x00\x00"
    files = {"file": ("programa.exe", exe_content, "application/octet-stream")}
    data = {"message": "Analiza este ejecutable"}

    response = client.post("/api/chat/document", data=data, files=files)

    assert response.status_code == 415
    assert "Unsupported file format" in response.json()["detail"]


def test_chat_document_fake_pdf_signature(client):
    """File has .pdf extension but invalid signature."""
    files = {"file": ("documento.pdf", b"This is not a PDF", "application/pdf")}
    data = {"message": "Analiza esto"}

    response = client.post("/api/chat/document", data=data, files=files)

    assert response.status_code == 415
    assert "missing PDF signature" in response.json()["detail"]


# ---------------------------------------------------------------------------
# 5. Archivo demasiado grande (413)
# ---------------------------------------------------------------------------
def test_chat_document_file_too_large(client):
    # Temporarily set max_file_size_mb to 1MB for this test to avoid allocating 10MB
    with patch.object(settings, "max_file_size_mb", 1):
        huge_content = b"A" * (1 * 1024 * 1024 + 100)
        files = {"file": ("grande.txt", huge_content, "text/plain")}
        data = {"message": "Archivo grande"}

        response = client.post("/api/chat/document", data=data, files=files)

        assert response.status_code == 413
        assert "exceeds maximum allowed size" in response.json()["detail"]


# ---------------------------------------------------------------------------
# 6. Archivo vacío (422)
# ---------------------------------------------------------------------------
def test_chat_document_empty_file(client):
    files = {"file": ("vacio.txt", b"", "text/plain")}
    data = {"message": "Archivo vacio"}

    response = client.post("/api/chat/document", data=data, files=files)

    assert response.status_code == 422
    assert "empty" in response.json()["detail"].lower()


def test_chat_document_whitespace_only_file(client):
    files = {"file": ("espacios.txt", b"    \n\n   \t  ", "text/plain")}
    data = {"message": "Archivo solo espacios"}

    response = client.post("/api/chat/document", data=data, files=files)

    assert response.status_code == 422
    assert "whitespace" in response.json()["detail"].lower()


# ---------------------------------------------------------------------------
# 7. /api/chat sigue funcionando intacto
# ---------------------------------------------------------------------------
def test_original_chat_endpoint_remains_functional(client, mock_openrouter):
    response = client.post("/api/chat", json={"message": "Hola, que servicios ofrece RCKT?"})

    assert response.status_code == 200
    json_data = response.json()
    assert "answer" in json_data
    assert json_data["answer"] == "Esta es una respuesta simulada del modelo."

    # Verify OpenRouter was called with original system prompt without document context
    assert mock_openrouter.called
    call_args = mock_openrouter.call_args[1]
    assert "## Contexto del documento adjunto" not in call_args["system"]
    assert call_args["user_message"] == "Hola, que servicios ofrece RCKT?"


# ---------------------------------------------------------------------------
# 8. /api/chat/document devuelve respuesta mockeada
# ---------------------------------------------------------------------------
def test_chat_document_mocked_response(client, mock_openrouter):
    files = {"file": ("reporte.txt", b"Ventas Q1: +20%", "text/plain")}
    data = {"message": "Cual fue el crecimiento?"}

    response = client.post("/api/chat/document", data=data, files=files)

    assert response.status_code == 200
    assert response.json() == {"answer": "Esta es una respuesta simulada del modelo."}


# ---------------------------------------------------------------------------
# 9. additional_context llega al Agent
# ---------------------------------------------------------------------------
def test_additional_context_passed_to_agent(client, mock_openrouter):
    doc_content = b"Informacion confidencial del proyecto Alpha: codigo 9942."
    files = {"file": ("alpha.txt", doc_content, "text/plain")}
    user_question = "Cual es el codigo del proyecto Alpha?"

    response = client.post("/api/chat/document", data={"message": user_question}, files=files)

    assert response.status_code == 200
    assert mock_openrouter.called
    call_args = mock_openrouter.call_args[1]

    # Verify temporal document context was injected into the system prompt
    assert "## Contexto del documento adjunto" in call_args["system"]
    assert "Informacion confidencial del proyecto Alpha: codigo 9942." in call_args["system"]
    assert call_args["user_message"] == user_question


# ---------------------------------------------------------------------------
# 10. Errores no exponen detalles internos
# ---------------------------------------------------------------------------
def test_errors_do_not_leak_internal_stack_traces(client):
    """When an unhandled exception or service failure occurs, no stack trace or secret is leaked."""
    with patch(
        "app.agent.agent.OpenRouterClient.complete_with_system",
        new_callable=AsyncMock,
        side_effect=Exception("OpenRouter API error: Connection timed out to https://openrouter.ai"),
    ):
        files = {"file": ("doc.txt", b"Texto normal", "text/plain")}
        data = {"message": "Pregunta"}

        response = client.post("/api/chat/document", data=data, files=files)

        assert response.status_code == 503
        data = response.json()
        assert "AI service unavailable" in data["detail"]
        # Ensure stack trace / tracebacks are not leaked in detail
        assert "Traceback" not in data["detail"]
        assert "openrouter_api_key" not in data["detail"].lower()


# ---------------------------------------------------------------------------
# PDF y DOCX válidos y sin texto (Casos complementarios)
# ---------------------------------------------------------------------------
def test_chat_document_valid_pdf(client, mock_openrouter):
    pdf_bytes = create_in_memory_pdf("Texto clave en PDF para el modelo.")
    files = {"file": ("informe.pdf", pdf_bytes, "application/pdf")}
    data = {"message": "Extrae el texto"}

    response = client.post("/api/chat/document", data=data, files=files)

    assert response.status_code == 200
    assert response.json()["answer"] == "Esta es una respuesta simulada del modelo."


def test_chat_document_scanned_pdf_without_text(client):
    """Scanned/blank PDF with no extractable text returns 422 with clear message."""
    empty_pdf_bytes = create_in_memory_empty_pdf()
    files = {"file": ("escaneado.pdf", empty_pdf_bytes, "application/pdf")}
    data = {"message": "Que dice?"}

    response = client.post("/api/chat/document", data=data, files=files)

    assert response.status_code == 422
    assert "scanned" in response.json()["detail"].lower() or "no extractable text" in response.json()["detail"].lower()


def test_chat_document_valid_docx(client, mock_openrouter):
    docx_bytes = create_in_memory_docx("Texto clave en DOCX para el asistente.")
    files = {"file": ("resumen.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    data = {"message": "Resume este DOCX"}

    response = client.post("/api/chat/document", data=data, files=files)

    assert response.status_code == 200
    assert response.json()["answer"] == "Esta es una respuesta simulada del modelo."


def test_chat_document_empty_docx(client):
    """Empty DOCX document returns 422."""
    empty_docx_bytes = create_in_memory_empty_docx()
    files = {"file": ("vacio.docx", empty_docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    data = {"message": "Que dice?"}

    response = client.post("/api/chat/document", data=data, files=files)

    assert response.status_code == 422
    assert "no extractable text" in response.json()["detail"].lower()


# ---------------------------------------------------------------------------
# Truncado de documento largo (> DOCUMENT_MAX_CHARS)
# ---------------------------------------------------------------------------
def test_document_truncation_indication():
    large_text = "A" * 55000
    parsed = parse_document("documento_largo.txt", large_text.encode("utf-8"))

    assert len(parsed) > 50000
    assert "[NOTA INTERNA: El documento superó el límite de 50000 caracteres" in parsed
